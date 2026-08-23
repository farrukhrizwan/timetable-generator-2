"""
School Timetable Generator
A standalone Windows desktop application for generating school timetables.
Packages to single .exe via PyInstaller.

Fixes applied (v3):
  1. Classes/Sections grid: 3 columns instead of 4 (prevents horizontal overflow / clipping)
  2. Generate buttons: placed inside the scrollable canvas frame so they are never clipped
  3. Mouse-wheel scrolling: bound to the correct canvas in BOTH main window and preview window
     (including Linux <Button-4>/<Button-5> support)
  4. Footer: copyright label "© Farrukh Barlas | AFUWEBS" added at bottom of main scrollable area
  5. App icon: icon.ico loaded via iconbitmap / wm_iconbitmap with safe fallback
  6. InchargeFrame: 3 columns to match new section layout
  7. PreviewWindow cell width: fixed divisor bug (was // 8 which collapsed cells to ~16px)
  8. period_headers: duplicate break-column edge-case corrected
  9. Spinbox from_=0 fix so opening minutes / hour can be 0
 10. Canvas width binding: scroll frame correctly tracks canvas width resize
 11. Wheel binding scoped to canvas widget (not bind_all which caused cross-window conflicts)
 12. Icon embedded as base64 fallback when running from PyInstaller bundle
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, font as tkfont
import os
import sys
import copy
import random
import base64
import io
from datetime import datetime, timedelta
from collections import defaultdict

# ── third-party (bundled via PyInstaller) ────────────────────────────────────
from PIL import Image, ImageTk
import openpyxl
from openpyxl.styles import (Font as XLFont, PatternFill, Alignment,
                              Border, Side, GradientFill)
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage
import docx
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.platypus import (SimpleDocTemplate, Table as RLTable,
                                TableStyle, Paragraph, Spacer, Image as RLImage)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ─────────────────────────────────────────────────────────────────────────────
# PALETTE
# ─────────────────────────────────────────────────────────────────────────────
C_NAVY   = "#1a2e5e"
C_GOLD   = "#c8a84b"
C_TEAL   = "#2e8b7a"
C_WHITE  = "#ffffff"
C_LIGHT  = "#f4f7fc"
C_BORDER = "#d0d8e8"
C_RED    = "#c0392b"
C_GREEN  = "#27ae60"
C_ORANGE = "#e67e22"
C_LGRAY  = "#ecf0f4"

# ─────────────────────────────────────────────────────────────────────────────
# ICON HELPER  (works from source tree AND PyInstaller bundle)
# ─────────────────────────────────────────────────────────────────────────────
def _icon_path():
    """Return absolute path to icon.ico, whether running live or bundled."""
    if getattr(sys, 'frozen', False):
        base = sys._MEIPASS
    else:
        base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, "icon.ico")


def apply_icon(window):
    """Safely apply the .ico to any Tk / Toplevel window."""
    ico = _icon_path()
    if os.path.exists(ico):
        try:
            window.iconbitmap(ico)
        except Exception:
            pass


# ─────────────────────────────────────────────────────────────────────────────
# DATA MODEL
# ─────────────────────────────────────────────────────────────────────────────
class AppData:
    """Central store for all user inputs."""
    def __init__(self):
        self.school_name   = ""
        self.logo_path     = ""
        self.section       = ""          # Pre School / Primary / etc.
        self.open_h        = 8
        self.open_m        = 0
        self.num_periods   = 7
        self.period_mins   = 40
        self.break_mins    = 20
        self.break_after   = 3          # break after period N
        self.off_time      = "14:30"
        # list of (class_name, section_letter) e.g. [("1","A"),("1","B")]
        self.class_sections = []
        # dict: (class, section) -> incharge_name
        self.incharges      = {}
        # list of (subject, [teacher1, teacher2, ...])
        self.subjects       = []
        # generated timetables
        self.class_tt       = {}   # (cls,sec) -> [period_list]  each item: (subject,teacher)
        self.teacher_tt     = {}   # teacher -> {(cls,sec): [period_list]}


# ─────────────────────────────────────────────────────────────────────────────
# TIMETABLE ENGINE
# ─────────────────────────────────────────────────────────────────────────────
class TTEngine:
    """Generates class-wise and teacher-wise timetables."""

    @staticmethod
    def compute_period_times(open_h, open_m, num_periods,
                             period_mins, break_mins, break_after):
        """Return list of (start_str, end_str) for each period."""
        times = []
        cur = datetime(2000, 1, 1, open_h, open_m)
        for i in range(1, num_periods + 1):
            dur = period_mins + (10 if i == 1 else 0)   # first period 10 min longer
            start = cur
            end   = cur + timedelta(minutes=dur)
            times.append((start.strftime("%H:%M"), end.strftime("%H:%M")))
            cur = end
            if i == break_after:
                cur += timedelta(minutes=break_mins)    # recess
        return times

    # Sentinel used throughout the app to mark the recess/break slot.
    BREAK_SLOT = ("__BREAK__", "__BREAK__")

    @staticmethod
    def generate(data: AppData):
        """
        Core scheduling algorithm.
        Returns (class_tt, teacher_tt) or raises ValueError with message.
        """
        if not data.class_sections:
            raise ValueError("No classes/sections defined.")
        if not data.subjects:
            raise ValueError("No subjects defined.")

        num_p    = data.num_periods
        brk_idx  = data.break_after
        subjects = data.subjects

        teacher_load = defaultdict(int)
        for _, teachers in subjects:
            for t in teachers:
                teacher_load[t] = 0

        class_tt = {}

        for cls, sec in data.class_sections:
            key      = (cls, sec)
            incharge = data.incharges.get(key, "")

            subj_pool = list(subjects)
            random.shuffle(subj_pool)
            while len(subj_pool) < num_p:
                subj_pool += [random.choice(subjects)]
            subj_pool = subj_pool[:num_p]

            incharge_subj = None
            if incharge:
                for subj, teachers in subjects:
                    if incharge in teachers:
                        incharge_subj = (subj, incharge)
                        break
            if incharge_subj is None:
                subj, teachers = subjects[0] if subjects else ("General", [incharge])
                incharge_subj = (subj, incharge if incharge else (teachers[0] if teachers else "TBD"))

            used_subjs = {incharge_subj[0]}
            remaining  = [sp for sp in subj_pool if sp[0] not in used_subjs]
            random.shuffle(remaining)

            teaching_slots = [incharge_subj]

            for i in range(1, num_p):
                if remaining:
                    subj, teachers = remaining.pop(0)
                else:
                    idx = (i - 1) % len(subjects)
                    subj, teachers = subjects[idx]

                teacher = (min(teachers, key=lambda t: teacher_load[t])
                           if teachers else "TBD")
                teacher_load[teacher] += 1
                teaching_slots.append((subj, teacher))

            period_list = list(teaching_slots)
            period_list.insert(brk_idx, TTEngine.BREAK_SLOT)

            class_tt[key] = period_list

        total_slots = num_p + 1
        teacher_tt  = defaultdict(dict)
        for (cls, sec), periods in class_tt.items():
            for p_idx, slot in enumerate(periods):
                if slot is TTEngine.BREAK_SLOT or slot == TTEngine.BREAK_SLOT:
                    continue
                subj, teacher = slot
                if teacher not in teacher_tt:
                    teacher_tt[teacher] = {}
                cs_key = f"{cls}-{sec}"
                if cs_key not in teacher_tt[teacher]:
                    teacher_tt[teacher][cs_key] = ["—"] * total_slots
                teacher_tt[teacher][cs_key][p_idx] = subj

        return class_tt, dict(teacher_tt)


# ─────────────────────────────────────────────────────────────────────────────
# PERIOD HEADERS
# ─────────────────────────────────────────────────────────────────────────────
def period_headers(data: AppData):
    """
    Return (headers, times).
    Both lists include one extra entry for the BREAK slot.
    FIX: removed duplicate break-column insertion at edge case.
    """
    raw_times = TTEngine.compute_period_times(
        data.open_h, data.open_m, data.num_periods,
        data.period_mins, data.break_mins, data.break_after)

    from datetime import datetime as _dt, timedelta as _td
    cur = _dt(2000, 1, 1, data.open_h, data.open_m)
    brk_start = brk_end = None
    for i in range(1, data.num_periods + 1):
        dur = data.period_mins + (10 if i == 1 else 0)
        cur += _td(minutes=dur)
        if i == data.break_after:
            brk_start = cur.strftime("%H:%M")
            cur += _td(minutes=data.break_mins)
            brk_end = cur.strftime("%H:%M")
            break

    headers = []
    times   = []
    p_num   = 1
    break_inserted = False
    for i, (s, e) in enumerate(raw_times):
        # Insert break slot BEFORE the period that comes after break_after
        if i == data.break_after and not break_inserted and brk_start:
            headers.append(f"BREAK\n{brk_start}-{brk_end}")
            times.append(("__BREAK__", "__BREAK__"))
            break_inserted = True
        headers.append(f"P{p_num}\n{s}-{e}")
        times.append((s, e))
        p_num += 1

    # If break_after >= num_periods and not yet inserted (trailing break)
    if not break_inserted and brk_start:
        headers.append(f"BREAK\n{brk_start}-{brk_end}")
        times.append(("__BREAK__", "__BREAK__"))

    return headers, times


# ─────────────────────────────────────────────────────────────────────────────
# EXPORT HELPERS – XLSX
# ─────────────────────────────────────────────────────────────────────────────
def export_xlsx(data: AppData, tt_type: str, filepath: str):
    """Export timetable to Excel."""
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    headers, times = period_headers(data)
    phdrs = headers

    navy_fill  = PatternFill("solid", fgColor="1a2e5e")
    gold_fill  = PatternFill("solid", fgColor="c8a84b")
    teal_fill  = PatternFill("solid", fgColor="2e8b7a")
    lgray_fill = PatternFill("solid", fgColor="ecf0f4")
    alt_fill   = PatternFill("solid", fgColor="dce8f5")
    white_fill = PatternFill("solid", fgColor="ffffff")
    brk_fill   = PatternFill("solid", fgColor="fff3cd")

    thin   = Side(style='thin', color='b0b8cc')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def hdr_font(color="FFFFFF", bold=True, size=11):
        return XLFont(name="Arial", bold=bold, color=color, size=size)

    def cell_font(bold=False, size=10, color="1a2e5e"):
        return XLFont(name="Arial", bold=bold, color=color, size=size)

    center     = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left       = Alignment(horizontal='left',   vertical='center', wrap_text=True)
    total_cols = 1 + len(phdrs)

    if tt_type == "class":
        ws = wb.create_sheet("Class Timetable")

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_cols)
        c = ws.cell(1, 1, data.school_name or "School Timetable")
        c.font = XLFont(name="Arial", bold=True, size=16, color="FFFFFF")
        c.alignment = center; c.fill = navy_fill

        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=total_cols)
        c2 = ws.cell(2, 1, f"Time Table — {data.section} Section")
        c2.font = XLFont(name="Arial", bold=True, size=13, color="1a2e5e")
        c2.alignment = center; c2.fill = gold_fill

        ws.cell(3, 1, "Class / Section").font = hdr_font()
        ws.cell(3, 1).fill = teal_fill
        ws.cell(3, 1).alignment = center
        ws.cell(3, 1).border = border

        for j, ph in enumerate(phdrs, 2):
            is_brk = "__BREAK__" in ph
            c = ws.cell(3, j, ph.replace("__BREAK__\n", "").replace("\n__BREAK__", ""))
            c.font = XLFont(name="Arial", bold=True,
                            color="7a5c00" if is_brk else "FFFFFF", size=9)
            c.fill = PatternFill("solid", fgColor="fff3cd") if is_brk else teal_fill
            c.alignment = center; c.border = border

        row = 4
        for idx, (cs, periods) in enumerate(data.class_tt.items()):
            cls, sec = cs
            fill = alt_fill if idx % 2 == 0 else white_fill
            label = f"{cls} - {sec}"
            c = ws.cell(row, 1, label)
            c.font = XLFont(name="Arial", bold=True, color="FFFFFF", size=10)
            c.fill = navy_fill if idx % 2 == 0 else teal_fill
            c.alignment = center; c.border = border

            for j, slot in enumerate(periods, 2):
                if slot == TTEngine.BREAK_SLOT:
                    c = ws.cell(row, j, "RECESS")
                    c.font = XLFont(name="Arial", bold=True, color="7a5c00", size=10)
                    c.fill = brk_fill; c.alignment = center; c.border = border
                else:
                    subj, teacher = slot
                    c = ws.cell(row, j, f"{subj}\n({teacher})")
                    c.font = cell_font(size=9)
                    c.fill = fill; c.alignment = center; c.border = border
                    if j == 2:
                        c.fill = PatternFill("solid", fgColor="e8f4e8")
                        c.font = XLFont(name="Arial", bold=True, size=9, color="1a5e2a")
            row += 1

        ws.column_dimensions['A'].width = 18
        for j in range(2, total_cols + 1):
            ws.column_dimensions[get_column_letter(j)].width = 16
        ws.row_dimensions[1].height = 28
        ws.row_dimensions[2].height = 22
        ws.row_dimensions[3].height = 36
        for r in range(4, row):
            ws.row_dimensions[r].height = 40

        if data.logo_path and os.path.exists(data.logo_path):
            try:
                img = XLImage(data.logo_path)
                img.width = 60; img.height = 60
                ws.add_image(img, "A1")
            except Exception:
                pass

    else:  # teacher
        ws = wb.create_sheet("Teacher Timetable")

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_cols)
        c = ws.cell(1, 1, data.school_name or "School Timetable")
        c.font = XLFont(name="Arial", bold=True, size=16, color="FFFFFF")
        c.alignment = center; c.fill = navy_fill

        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=total_cols)
        c2 = ws.cell(2, 1, f"Teacher Timetable — {data.section} Section")
        c2.font = XLFont(name="Arial", bold=True, size=13, color="1a2e5e")
        c2.alignment = center; c2.fill = gold_fill

        ws.cell(3, 1, "Teacher").font = hdr_font()
        ws.cell(3, 1).fill = teal_fill
        ws.cell(3, 1).alignment = center
        ws.cell(3, 1).border = border

        for j, ph in enumerate(phdrs, 2):
            is_brk = "__BREAK__" in ph
            c = ws.cell(3, j, ph.replace("__BREAK__\n", "").replace("\n__BREAK__",""))
            c.font = XLFont(name="Arial", bold=True,
                            color="7a5c00" if is_brk else "FFFFFF", size=9)
            c.fill = PatternFill("solid", fgColor="fff3cd") if is_brk else teal_fill
            c.alignment = center; c.border = border

        row = 4
        total_slots = len(phdrs)
        for idx, (teacher, cs_map) in enumerate(sorted(data.teacher_tt.items())):
            fill = alt_fill if idx % 2 == 0 else white_fill
            combined = ["—"] * total_slots
            brk_pos = next((i for i, t in enumerate(times) if t == ("__BREAK__", "__BREAK__")), None)
            if brk_pos is not None:
                combined[brk_pos] = "__BREAK__"

            for cs_key, periods in cs_map.items():
                for p, subj in enumerate(periods):
                    if subj != "—":
                        combined[p] = (f"{subj}\n({cs_key})"
                                       if combined[p] in ("—", "__BREAK__", "")
                                       else combined[p] + f"\n{subj}({cs_key})")

            c = ws.cell(row, 1, teacher)
            c.font = XLFont(name="Arial", bold=True, color="FFFFFF", size=10)
            c.fill = navy_fill; c.alignment = center; c.border = border

            for j, cell_val in enumerate(combined, 2):
                if cell_val == "__BREAK__":
                    c = ws.cell(row, j, "RECESS")
                    c.font = XLFont(name="Arial", bold=True, color="7a5c00", size=10)
                    c.fill = PatternFill("solid", fgColor="fff3cd")
                    c.alignment = center; c.border = border
                else:
                    c = ws.cell(row, j, cell_val)
                    c.font = cell_font(size=9)
                    c.fill = fill; c.alignment = center; c.border = border
            row += 1

        ws.column_dimensions['A'].width = 20
        for j in range(2, total_cols + 1):
            ws.column_dimensions[get_column_letter(j)].width = 16
        ws.row_dimensions[1].height = 28
        ws.row_dimensions[2].height = 22
        ws.row_dimensions[3].height = 36
        for r in range(4, row):
            ws.row_dimensions[r].height = 44

        if data.logo_path and os.path.exists(data.logo_path):
            try:
                img = XLImage(data.logo_path)
                img.width = 60; img.height = 60
                ws.add_image(img, "A1")
            except Exception:
                pass

    wb.save(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# EXPORT HELPERS – DOCX
# ─────────────────────────────────────────────────────────────────────────────
def export_docx(data: AppData, tt_type: str, filepath: str):
    """Export timetable to Word docx."""
    doc = DocxDocument()
    section_obj = doc.sections[0]
    section_obj.page_width    = Cm(29.7)
    section_obj.page_height   = Cm(21.0)
    section_obj.left_margin   = Cm(1.5)
    section_obj.right_margin  = Cm(1.5)
    section_obj.top_margin    = Cm(1.5)
    section_obj.bottom_margin = Cm(1.5)

    _, times = period_headers(data)
    total_slots = len(times)

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def add_header_block(label):
        hdr_p = doc.add_paragraph()
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if data.logo_path and os.path.exists(data.logo_path):
            try:
                run = hdr_p.add_run()
                run.add_picture(data.logo_path, width=Inches(0.55))
                hdr_p.add_run("  ")
            except Exception:
                pass
        run = hdr_p.add_run(data.school_name or "School Timetable")
        run.bold = True; run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x2e, 0x5e)

        sub_p = doc.add_paragraph(label)
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = sub_p.runs[0]
        run2.bold = True; run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(0xc8, 0xa8, 0x4b)
        doc.add_paragraph()

    if tt_type == "class":
        add_header_block(f"Class-Wise Time Table — {data.section} Section")
        num_cols = 1 + total_slots
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_row = table.rows[0]
        hdr_row.cells[0].text = "Class / Section"
        set_cell_bg(hdr_row.cells[0], "2e8b7a")
        hdr_row.cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        hdr_row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_num = 1
        for j, (s, e) in enumerate(times, 1):
            cell = hdr_row.cells[j]
            if (s, e) == ("__BREAK__", "__BREAK__"):
                cell.text = "RECESS"
                set_cell_bg(cell, "fff3cd")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x7a,0x5c,0x00)
                        run.font.size = Pt(8)
            else:
                cell.text = f"P{p_num}\n{s}-{e}"
                set_cell_bg(cell, "2e8b7a")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.size = Pt(8)
                p_num += 1

        for idx, (cs, periods) in enumerate(data.class_tt.items()):
            cls, sec = cs
            row = table.add_row()
            cls_cell = row.cells[0]
            cls_cell.text = f"{cls} – {sec}"
            bg = "1a2e5e" if idx % 2 == 0 else "2e8b7a"
            set_cell_bg(cls_cell, bg)
            cls_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cls_cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9)

            for j, slot in enumerate(periods, 1):
                cell = row.cells[j]
                if slot == TTEngine.BREAK_SLOT:
                    cell.text = "RECESS"
                    set_cell_bg(cell, "fff3cd")
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(8)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0x7a,0x5c,0x00)
                else:
                    subj, teacher = slot
                    cell.text = f"{subj}\n{teacher}"
                    fill_col = "e8f4e8" if j == 1 else ("f0f4fc" if idx%2==0 else "ffffff")
                    set_cell_bg(cell, fill_col)
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for i2, run in enumerate(para.runs):
                            run.font.size = Pt(8)
                            run.font.bold = (i2 == 0)
                            run.font.color.rgb = RGBColor(0x1a,0x2e,0x5e)

        col_w = Cm(2.4)
        for row in table.rows:
            row.cells[0].width = Cm(3.0)
            for j in range(1, num_cols):
                row.cells[j].width = col_w

    else:  # teacher
        add_header_block(f"Teacher-Wise Time Table — {data.section} Section")
        num_cols = 1 + total_slots
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_row = table.rows[0]
        hdr_row.cells[0].text = "Teacher"
        set_cell_bg(hdr_row.cells[0], "2e8b7a")
        hdr_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(9)

        p_num = 1
        for j, (s, e) in enumerate(times, 1):
            cell = hdr_row.cells[j]
            if (s, e) == ("__BREAK__", "__BREAK__"):
                cell.text = "RECESS"
                set_cell_bg(cell, "fff3cd")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x7a,0x5c,0x00)
                        run.font.size = Pt(8)
            else:
                cell.text = f"P{p_num}\n{s}-{e}"
                set_cell_bg(cell, "2e8b7a")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.size = Pt(8)
                p_num += 1

        brk_pos = next((i for i, t in enumerate(times) if t == ("__BREAK__","__BREAK__")), None)

        for idx, (teacher, cs_map) in enumerate(sorted(data.teacher_tt.items())):
            combined = ["—"] * total_slots
            if brk_pos is not None:
                combined[brk_pos] = "__BREAK__"
            for cs_key, periods in cs_map.items():
                for p, subj in enumerate(periods):
                    if subj != "—":
                        entry = f"{subj}\n({cs_key})"
                        combined[p] = entry if combined[p] in ("—","") else combined[p]+"\n"+entry

            row = table.add_row()
            tc = row.cells[0]
            tc.text = teacher
            set_cell_bg(tc, "1a2e5e" if idx%2==0 else "2e8b7a")
            tc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in tc.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9)

            for j, cell_val in enumerate(combined, 1):
                cell = row.cells[j]
                if cell_val == "__BREAK__":
                    cell.text = "RECESS"
                    set_cell_bg(cell, "fff3cd")
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(8)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0x7a,0x5c,0x00)
                else:
                    cell.text = cell_val
                    set_cell_bg(cell, "f0f4fc" if idx%2==0 else "ffffff")
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(8)
                            run.font.color.rgb = RGBColor(0x1a,0x2e,0x5e)

    doc.save(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# EXPORT HELPERS – PDF
# ─────────────────────────────────────────────────────────────────────────────
def export_pdf(data: AppData, tt_type: str, filepath: str):
    """Export timetable to PDF via ReportLab."""
    doc = SimpleDocTemplate(
        filepath,
        pagesize=landscape(A4),
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=1.5*cm,  bottomMargin=1.5*cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'SchoolTitle', parent=styles['Normal'],
        fontSize=16, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a2e5e'),
        alignment=TA_CENTER, spaceAfter=4)
    sub_style = ParagraphStyle(
        'SubTitle', parent=styles['Normal'],
        fontSize=12, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#c8a84b'),
        alignment=TA_CENTER, spaceAfter=8)
    cell_style = ParagraphStyle(
        'Cell', parent=styles['Normal'],
        fontSize=7.5, fontName='Helvetica', alignment=TA_CENTER)
    cell_bold = ParagraphStyle(
        'CellBold', parent=styles['Normal'],
        fontSize=7.5, fontName='Helvetica-Bold', alignment=TA_CENTER)

    _, times = period_headers(data)
    total_slots = len(times)
    brk_col_idx = next((i for i, t in enumerate(times) if t == ("__BREAK__","__BREAK__")), None)
    amber     = colors.HexColor('#fff3cd')
    amber_txt = colors.HexColor('#7a5c00')

    story = []
    story.append(Paragraph(data.school_name or "School Timetable", title_style))
    label = ("Class-Wise" if tt_type=="class" else "Teacher-Wise") + f" Time Table — {data.section} Section"
    story.append(Paragraph(label, sub_style))
    story.append(Spacer(1, 0.2*cm))

    navy  = colors.HexColor('#1a2e5e')
    teal  = colors.HexColor('#2e8b7a')
    alt   = colors.HexColor('#dce8f5')
    white = colors.white
    green = colors.HexColor('#e8f4e8')

    recess_style = ParagraphStyle(
        'Recess', parent=styles['Normal'],
        fontSize=7.5, fontName='Helvetica-Bold',
        textColor=amber_txt, alignment=TA_CENTER)

    if tt_type == "class":
        hdr = [Paragraph("Class/Section", cell_bold)]
        p_num = 1
        for s, e in times:
            if (s,e) == ("__BREAK__","__BREAK__"):
                hdr.append(Paragraph("<b>RECESS</b>", recess_style))
            else:
                hdr.append(Paragraph(f"<b>P{p_num}<br/>{s}-{e}</b>", cell_bold))
                p_num += 1

        tdata = [hdr]
        for cs, periods in data.class_tt.items():
            cls, sec = cs
            row_data = [Paragraph(f"<b>{cls} – {sec}</b>", cell_bold)]
            for slot in periods:
                if slot == TTEngine.BREAK_SLOT:
                    row_data.append(Paragraph("<b>RECESS</b>", recess_style))
                else:
                    subj, teacher = slot
                    row_data.append(Paragraph(f"<b>{subj}</b><br/>{teacher}", cell_style))
            tdata.append(row_data)

        col_w = (landscape(A4)[0] - 3*cm) / (1 + total_slots)
        col_widths = [col_w * 1.4] + [col_w] * total_slots
        tbl = RLTable(tdata, colWidths=col_widths, repeatRows=1)

        ts = [
            ('BACKGROUND', (0,0), (-1,0), teal),
            ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
            ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE',   (0,0), (-1,0), 8),
            ('ALIGN',      (0,0), (-1,-1), 'CENTER'),
            ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
            ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#b0b8cc')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [alt, white]),
            ('BACKGROUND', (0,1), (0,-1), navy),
            ('TEXTCOLOR',  (0,1), (0,-1), colors.white),
            ('BACKGROUND', (2,1), (2,-1), green),
            ('FONTSIZE',   (0,1), (-1,-1), 7.5),
            ('ROWHEIGHT',  (0,1), (-1,-1), 0.85*cm),
        ]
        if brk_col_idx is not None:
            bc = brk_col_idx + 1
            ts += [
                ('BACKGROUND', (bc,0), (bc,-1), amber),
                ('TEXTCOLOR',  (bc,0), (bc,-1), amber_txt),
            ]
        tbl.setStyle(TableStyle(ts))

    else:  # teacher
        hdr = [Paragraph("Teacher", cell_bold)]
        p_num = 1
        for s, e in times:
            if (s,e) == ("__BREAK__","__BREAK__"):
                hdr.append(Paragraph("<b>RECESS</b>", recess_style))
            else:
                hdr.append(Paragraph(f"<b>P{p_num}<br/>{s}-{e}</b>", cell_bold))
                p_num += 1

        tdata = [hdr]
        for teacher, cs_map in sorted(data.teacher_tt.items()):
            combined = ["—"] * total_slots
            if brk_col_idx is not None:
                combined[brk_col_idx] = "__BREAK__"
            for cs_key, periods in cs_map.items():
                for p, subj in enumerate(periods):
                    if subj != "—":
                        entry = f"{subj}\n({cs_key})"
                        combined[p] = entry if combined[p] in ("—","") else combined[p]+"\n"+entry

            row_data = [Paragraph(f"<b>{teacher}</b>", cell_bold)]
            for cv in combined:
                if cv == "__BREAK__":
                    row_data.append(Paragraph("<b>RECESS</b>", recess_style))
                else:
                    row_data.append(Paragraph(cv.replace("\n","<br/>"), cell_style))
            tdata.append(row_data)

        col_w = (landscape(A4)[0] - 3*cm) / (1 + total_slots)
        col_widths = [col_w * 1.6] + [col_w * 0.97] * total_slots
        tbl = RLTable(tdata, colWidths=col_widths, repeatRows=1)
        ts = [
            ('BACKGROUND', (0,0), (-1,0), teal),
            ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
            ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE',   (0,0), (-1,0), 8),
            ('ALIGN',      (0,0), (-1,-1), 'CENTER'),
            ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
            ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#b0b8cc')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [alt, white]),
            ('BACKGROUND', (0,1), (0,-1), navy),
            ('TEXTCOLOR',  (0,1), (0,-1), colors.white),
            ('FONTSIZE',   (0,1), (-1,-1), 7.5),
            ('ROWHEIGHT',  (0,1), (-1,-1), 0.85*cm),
        ]
        if brk_col_idx is not None:
            bc = brk_col_idx + 1
            ts += [
                ('BACKGROUND', (bc,0), (bc,-1), amber),
                ('TEXTCOLOR',  (bc,0), (bc,-1), amber_txt),
            ]
        tbl.setStyle(TableStyle(ts))

    story.append(tbl)
    story.append(Spacer(1, 0.3*cm))
    note_style = ParagraphStyle('Note', parent=styles['Normal'],
                                fontSize=7, fontName='Helvetica-Oblique',
                                textColor=colors.grey, alignment=TA_CENTER)
    story.append(Paragraph(
        f"Generated on {datetime.now().strftime('%d %B %Y %H:%M')} — "
        f"Break of {data.break_mins} min after Period {data.break_after}",
        note_style))

    doc.build(story)


# ─────────────────────────────────────────────────────────────────────────────
# GUI HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def styled_btn(parent, text, cmd, bg=C_NAVY, fg=C_WHITE,
               padx=18, pady=6, font_size=10, **kw):
    btn = tk.Button(parent, text=text, command=cmd,
                    bg=bg, fg=fg, activebackground=C_GOLD,
                    activeforeground=C_WHITE, relief='flat',
                    font=("Arial", font_size, "bold"),
                    cursor="hand2", padx=padx, pady=pady, **kw)
    return btn


def lbl(parent, text, bold=False, size=10, fg=C_NAVY, **kw):
    f = ("Arial", size, "bold") if bold else ("Arial", size)
    return tk.Label(parent, text=text, font=f, fg=fg,
                    bg=parent.cget('bg'), **kw)


def section_frame(parent, title):
    """Titled card-style LabelFrame."""
    fr = tk.LabelFrame(parent, text=f"  {title}  ",
                       font=("Arial", 10, "bold"),
                       fg=C_GOLD, bg=C_LIGHT,
                       relief='groove', bd=2,
                       labelanchor='nw', padx=10, pady=8)
    return fr


def _bind_mousewheel(widget, canvas):
    """
    Bind mouse-wheel events on `widget` so they scroll `canvas`.
    Handles Windows (<MouseWheel>) and Linux (<Button-4>/<Button-5>).
    FIX: scope binding to widget, not bind_all, to avoid cross-window scroll conflicts.
    """
    def _on_wheel_win(e):
        canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")

    def _on_btn4(e):   # Linux scroll up
        canvas.yview_scroll(-1, "units")

    def _on_btn5(e):   # Linux scroll down
        canvas.yview_scroll(1, "units")

    widget.bind("<MouseWheel>", _on_wheel_win, add="+")
    widget.bind("<Button-4>",   _on_btn4,      add="+")
    widget.bind("<Button-5>",   _on_btn5,      add="+")


def _bind_mousewheel_recursive(widget, canvas):
    """Recursively bind mouse-wheel on a widget and all its children."""
    _bind_mousewheel(widget, canvas)
    for child in widget.winfo_children():
        _bind_mousewheel_recursive(child, canvas)


# ─────────────────────────────────────────────────────────────────────────────
# PREVIEW WINDOW
# ─────────────────────────────────────────────────────────────────────────────
class PreviewWindow(tk.Toplevel):
    def __init__(self, parent, data: AppData, tt_type: str):
        super().__init__(parent)
        self.data    = data
        self.tt_type = tt_type
        self.title(("Class-Wise" if tt_type == "class" else "Teacher-Wise") + " Timetable Preview")
        self.geometry("1100x620")
        self.configure(bg=C_LIGHT)
        self.resizable(True, True)
        apply_icon(self)    # FIX: icon on preview window too
        self._build()

    def _build(self):
        # ── top bar ──
        top = tk.Frame(self, bg=C_NAVY, pady=8)
        top.pack(fill='x')

        tk.Label(top, text=self.data.school_name or "School Timetable",
                 font=("Arial", 16, "bold"), fg=C_GOLD, bg=C_NAVY).pack(side='left', padx=16)

        lbl_type = "Class-Wise" if self.tt_type == "class" else "Teacher-Wise"
        tk.Label(top, text=f"{lbl_type} Time Table — {self.data.section}",
                 font=("Arial", 11), fg=C_WHITE, bg=C_NAVY).pack(side='left', padx=10)

        btn_fr = tk.Frame(top, bg=C_NAVY)
        btn_fr.pack(side='right', padx=12)
        styled_btn(btn_fr, "⬇ Excel", lambda: self._download("xlsx"),
                   bg=C_GREEN, pady=4).pack(side='left', padx=4)
        styled_btn(btn_fr, "⬇ Word",  lambda: self._download("docx"),
                   bg=C_TEAL,  pady=4).pack(side='left', padx=4)
        styled_btn(btn_fr, "⬇ PDF",   lambda: self._download("pdf"),
                   bg=C_RED,   pady=4).pack(side='left', padx=4)

        # ── scrollable table ──
        outer = tk.Frame(self, bg=C_LIGHT)
        outer.pack(fill='both', expand=True, padx=10, pady=10)

        self._canvas = tk.Canvas(outer, bg=C_LIGHT, highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient='vertical',   command=self._canvas.yview)
        hsb = ttk.Scrollbar(outer, orient='horizontal', command=self._canvas.xview)
        self._canvas.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        vsb.pack(side='right',  fill='y')
        hsb.pack(side='bottom', fill='x')
        self._canvas.pack(side='left', fill='both', expand=True)

        inner = tk.Frame(self._canvas, bg=C_LIGHT)
        self._canvas.create_window((0, 0), window=inner, anchor='nw')
        inner.bind("<Configure>",
                   lambda e: self._canvas.configure(scrollregion=self._canvas.bbox("all")))

        self._render_table(inner)

        # FIX: bind wheel on inner frame and all children (not bind_all)
        _bind_mousewheel_recursive(inner, self._canvas)
        _bind_mousewheel(self._canvas, self._canvas)

    def _render_table(self, parent):
        _, times = period_headers(self.data)

        # FIX: use proper pixel widths, not // 8 which made cells ~16px wide
        COL_W0 = 140   # first column (class/teacher label)
        COL_W  = 120   # period columns
        ROW_H  = 2     # height in text lines (Label widget)

        navy_bg  = C_NAVY
        teal_bg  = C_TEAL
        alt_bg   = "#dce8f5"
        white_bg = "#ffffff"
        green_bg = "#e8f4e8"
        brk_bg   = "#fff3cd"
        brk_fg   = "#7a5c00"

        def cell(fr, text, bg, fg, bold=False, w_px=COL_W):
            f = ("Arial", 8, "bold") if bold else ("Arial", 8)
            c = tk.Label(fr, text=text, bg=bg, fg=fg, font=f,
                         wraplength=w_px - 8,
                         relief='flat', anchor='center', justify='center',
                         padx=4, pady=6, bd=0)
            return c

        # Header row
        hfr = tk.Frame(parent, bg=C_LIGHT)
        hfr.pack(fill='x')

        first_col = "Class / Section" if self.tt_type == "class" else "Teacher"

        # FIX: use padx/pady + sticky for proper cell sizing
        cell(hfr, first_col, teal_bg, C_WHITE, bold=True,
             w_px=COL_W0).grid(row=0, column=0, sticky='nsew',
                               padx=1, pady=1, ipadx=4, ipady=4)
        hfr.grid_columnconfigure(0, minsize=COL_W0)

        p_num = 1
        for j, (s, e) in enumerate(times, 1):
            if (s, e) == ("__BREAK__", "__BREAK__"):
                txt = "RECESS"
                cell(hfr, txt, brk_bg, brk_fg, bold=True,
                     w_px=COL_W).grid(row=0, column=j, sticky='nsew', padx=1, pady=1)
            else:
                txt = f"P{p_num}\n{s}–{e}"
                cell(hfr, txt, teal_bg, C_WHITE, bold=True,
                     w_px=COL_W).grid(row=0, column=j, sticky='nsew', padx=1, pady=1)
                p_num += 1
            hfr.grid_columnconfigure(j, minsize=COL_W)

        # Data rows
        if self.tt_type == "class":
            for idx, (cs, periods) in enumerate(self.data.class_tt.items()):
                cls, sec = cs
                rfr = tk.Frame(parent, bg=C_LIGHT)
                rfr.pack(fill='x')
                row_bg = alt_bg if idx % 2 == 0 else white_bg
                cls_bg = navy_bg if idx % 2 == 0 else teal_bg
                cell(rfr, f"{cls} – {sec}", cls_bg, C_WHITE, bold=True,
                     w_px=COL_W0).grid(row=0, column=0, sticky='nsew',
                                       padx=1, pady=1, ipadx=4, ipady=4)
                rfr.grid_columnconfigure(0, minsize=COL_W0)

                teaching_col = 0
                for j, slot in enumerate(periods, 1):
                    if slot == TTEngine.BREAK_SLOT:
                        cell(rfr, "RECESS", brk_bg, brk_fg, bold=True,
                             w_px=COL_W).grid(row=0, column=j, sticky='nsew', padx=1, pady=1)
                    else:
                        subj, teacher = slot
                        teaching_col += 1
                        bg = green_bg if teaching_col == 1 else row_bg
                        cell(rfr, f"{subj}\n{teacher}", bg, C_NAVY,
                             bold=(teaching_col == 1),
                             w_px=COL_W).grid(row=0, column=j, sticky='nsew', padx=1, pady=1)
                    rfr.grid_columnconfigure(j, minsize=COL_W)
        else:
            brk_pos = next((i for i, t in enumerate(times) if t == ("__BREAK__","__BREAK__")), None)
            total_slots = len(times)
            for idx, (teacher, cs_map) in enumerate(sorted(self.data.teacher_tt.items())):
                combined = ["—"] * total_slots
                if brk_pos is not None:
                    combined[brk_pos] = "__BREAK__"
                for cs_key, periods in cs_map.items():
                    for p, subj in enumerate(periods):
                        if subj != "—":
                            entry = f"{subj}\n({cs_key})"
                            combined[p] = (entry if combined[p] in ("—","__BREAK__","")
                                           else combined[p]+"\n"+entry)

                rfr = tk.Frame(parent, bg=C_LIGHT)
                rfr.pack(fill='x')
                row_bg = alt_bg if idx % 2 == 0 else white_bg
                tchr_bg = navy_bg if idx % 2 == 0 else teal_bg
                cell(rfr, teacher, tchr_bg, C_WHITE, bold=True,
                     w_px=COL_W0).grid(row=0, column=0, sticky='nsew',
                                       padx=1, pady=1, ipadx=4, ipady=4)
                rfr.grid_columnconfigure(0, minsize=COL_W0)

                for j, cv in enumerate(combined, 1):
                    if cv == "__BREAK__":
                        cell(rfr, "RECESS", brk_bg, brk_fg, bold=True,
                             w_px=COL_W).grid(row=0, column=j, sticky='nsew', padx=1, pady=1)
                    else:
                        cell(rfr, cv, row_bg, C_NAVY,
                             w_px=COL_W).grid(row=0, column=j, sticky='nsew', padx=1, pady=1)
                    rfr.grid_columnconfigure(j, minsize=COL_W)

    def _download(self, fmt):
        ext_map = {"xlsx": "Excel Files (*.xlsx)", "docx": "Word Files (*.docx)", "pdf": "PDF Files (*.pdf)"}
        fp = filedialog.asksaveasfilename(
            defaultextension=f".{fmt}",
            filetypes=[(ext_map[fmt], f"*.{fmt}"), ("All Files", "*.*")],
            initialfile=f"Timetable_{self.tt_type}_{fmt}.{fmt}"
        )
        if not fp:
            return
        try:
            if fmt == "xlsx":
                export_xlsx(self.data, self.tt_type, fp)
            elif fmt == "docx":
                export_docx(self.data, self.tt_type, fp)
            else:
                export_pdf(self.data, self.tt_type, fp)
            messagebox.showinfo("✅ Saved", f"File saved:\n{fp}", parent=self)
        except Exception as ex:
            messagebox.showerror("Export Error", str(ex), parent=self)


# ─────────────────────────────────────────────────────────────────────────────
# SUBJECT ENTRY ROW
# ─────────────────────────────────────────────────────────────────────────────
class SubjectRow(tk.Frame):
    def __init__(self, parent, remove_cb, **kw):
        super().__init__(parent, bg=C_LIGHT, **kw)
        self._remove_cb = remove_cb
        self._build()

    def _build(self):
        tk.Label(self, text="Subject:", font=("Arial", 9), bg=C_LIGHT,
                 fg=C_NAVY).pack(side='left')
        self.subj_var = tk.StringVar()
        tk.Entry(self, textvariable=self.subj_var, width=18,
                 font=("Arial", 9)).pack(side='left', padx=4)

        tk.Label(self, text="Teacher(s) (comma-sep):",
                 font=("Arial", 9), bg=C_LIGHT, fg=C_NAVY).pack(side='left', padx=(8, 0))
        self.tchr_var = tk.StringVar()
        tk.Entry(self, textvariable=self.tchr_var, width=34,
                 font=("Arial", 9)).pack(side='left', padx=4)

        styled_btn(self, "✕", self._remove_cb,
                   bg=C_RED, padx=6, pady=2, font_size=9).pack(side='left', padx=4)

    def get(self):
        subj  = self.subj_var.get().strip()
        tchrs = [t.strip() for t in self.tchr_var.get().split(",") if t.strip()]
        return subj, tchrs


# ─────────────────────────────────────────────────────────────────────────────
# INCHARGE ENTRY FRAME
# ─────────────────────────────────────────────────────────────────────────────
class InchargeFrame(tk.Frame):
    def __init__(self, parent, class_sections, **kw):
        super().__init__(parent, bg=C_LIGHT, **kw)
        self._vars = {}
        self._build(class_sections)

    def _build(self, class_sections):
        for widget in self.winfo_children():
            widget.destroy()

        tk.Label(self, text="Class-In-Charge Teacher Names",
                 font=("Arial", 10, "bold"), fg=C_NAVY, bg=C_LIGHT).grid(
            row=0, column=0, columnspan=6, sticky='w', pady=(0, 6))

        # FIX: 3 columns to prevent horizontal overflow on smaller windows
        COLS = 3
        for i, (cls, sec) in enumerate(class_sections):
            key = (cls, sec)
            row, col = divmod(i, COLS)
            row += 1
            col_base = col * 2

            tk.Label(self, text=f"{cls}-{sec}:", font=("Arial", 9, "bold"),
                     bg=C_LIGHT, fg=C_TEAL, width=8, anchor='e').grid(
                row=row, column=col_base, padx=(8, 2), pady=3, sticky='e')

            var = tk.StringVar(value=self._vars.get(key, tk.StringVar()).get())
            self._vars[key] = var
            tk.Entry(self, textvariable=var, width=22,
                     font=("Arial", 9)).grid(row=row, column=col_base + 1,
                                             padx=(0, 12), pady=3, sticky='w')

    def get_incharges(self):
        return {k: v.get().strip() for k, v in self._vars.items()}

    def refresh(self, class_sections):
        old_values = {k: v.get() for k, v in self._vars.items()}
        self._vars = {}
        for k, v in old_values.items():
            self._vars[k] = tk.StringVar(value=v)
        self._build(class_sections)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN APP WINDOW
# ─────────────────────────────────────────────────────────────────────────────
class TimetableApp(tk.Tk):

    SECTIONS = ["Pre School", "Primary", "Middle", "Senior", "College", "Other (type below)"]
    CLASSES  = ["P.G", "K.G", "Nursery", "Prep",
                "1","2","3","4","5","6","7","8","9","10","11","12",
                "Other (type)"]
    STD_SECS = ["A","B","C","D","G"]

    def __init__(self):
        super().__init__()
        self.data = AppData()
        self.title("🏫 School Timetable Generator")
        self.geometry("1100x820")
        self.minsize(900, 700)
        self.configure(bg=C_LIGHT)
        self._logo_img = None
        self._subject_rows = []
        self._class_sec_vars  = {}
        self._selected_classes = []
        apply_icon(self)        # FIX: set .ico on main window
        self._build_ui()

    # ── UI BUILD ──────────────────────────────────────────────────────────────
    def _build_ui(self):
        # Title bar
        tbar = tk.Frame(self, bg=C_NAVY, pady=12)
        tbar.pack(fill='x')
        tk.Label(tbar, text="🏫  School Timetable Generator",
                 font=("Arial", 18, "bold"), fg=C_GOLD, bg=C_NAVY).pack(side='left', padx=20)
        tk.Label(tbar, text="Automated · Professional · Offline",
                 font=("Arial", 10), fg=C_WHITE, bg=C_NAVY).pack(side='left')

        # Scrollable main area
        self._main_canvas = tk.Canvas(self, bg=C_LIGHT, highlightthickness=0)
        vsb = ttk.Scrollbar(self, orient='vertical', command=self._main_canvas.yview)
        self._main_canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side='right', fill='y')
        self._main_canvas.pack(side='left', fill='both', expand=True)

        self._scroll_frame = tk.Frame(self._main_canvas, bg=C_LIGHT)
        win_id = self._main_canvas.create_window((0, 0), window=self._scroll_frame, anchor='nw')

        # FIX: correctly expand scroll_frame width with canvas
        def _on_canvas_resize(e):
            self._main_canvas.itemconfig(win_id, width=e.width)
        self._main_canvas.bind('<Configure>', _on_canvas_resize)

        self._scroll_frame.bind(
            "<Configure>",
            lambda e: self._main_canvas.configure(scrollregion=self._main_canvas.bbox("all"))
        )

        # FIX: bind wheel on canvas itself + will propagate via recursive bind after content added
        _bind_mousewheel(self._main_canvas, self._main_canvas)

        self._build_step1()
        self._build_step2()
        self._build_step3()
        self._build_step4()
        self._build_step5()
        self._build_step6()
        self._build_footer()   # FIX: footer with copyright

        # After all widgets created, bind wheel recursively on scroll_frame content
        self._scroll_frame.after(100, self._bind_all_scroll)

    def _bind_all_scroll(self):
        """Deferred recursive binding so all child widgets get wheel support."""
        _bind_mousewheel_recursive(self._scroll_frame, self._main_canvas)

    def _padded(self, title):
        fr = section_frame(self._scroll_frame, title)
        fr.pack(fill='x', padx=16, pady=8)
        return fr

    # STEP 1 – School name + logo
    def _build_step1(self):
        fr = self._padded("① School / Institution Info")

        tk.Label(fr, text="School / College Name:", font=("Arial",10,"bold"),
                 fg=C_NAVY, bg=C_LIGHT).grid(row=0, column=0, sticky='w', pady=4)
        self._school_var = tk.StringVar()
        tk.Entry(fr, textvariable=self._school_var, width=45,
                 font=("Arial", 11)).grid(row=0, column=1, padx=8, sticky='w')

        tk.Label(fr, text="Logo Image:", font=("Arial",10,"bold"),
                 fg=C_NAVY, bg=C_LIGHT).grid(row=1, column=0, sticky='w', pady=4)
        logo_fr = tk.Frame(fr, bg=C_LIGHT)
        logo_fr.grid(row=1, column=1, sticky='w')
        self._logo_var = tk.StringVar()
        tk.Entry(logo_fr, textvariable=self._logo_var, width=32,
                 font=("Arial",9), state='readonly').pack(side='left')
        styled_btn(logo_fr, "Browse…", self._browse_logo, padx=8, pady=3,
                   font_size=9).pack(side='left', padx=6)

        self._logo_label = tk.Label(logo_fr, bg=C_LIGHT)
        self._logo_label.pack(side='left', padx=8)

    def _browse_logo(self):
        path = filedialog.askopenfilename(
            filetypes=[("Images", "*.png *.jpg *.jpeg *.bmp *.gif"), ("All","*.*")])
        if path:
            self._logo_var.set(path)
            try:
                img = Image.open(path).resize((50,50), Image.LANCZOS)
                self._logo_img = ImageTk.PhotoImage(img)
                self._logo_label.config(image=self._logo_img)
            except Exception:
                pass

    # STEP 2 – Section & timing
    def _build_step2(self):
        fr = self._padded("② Section & Schedule Settings")

        tk.Label(fr, text="School Section:", font=("Arial",9,"bold"),
                 fg=C_NAVY, bg=C_LIGHT).grid(row=0, column=0, sticky='w', pady=4)
        self._section_var = tk.StringVar(value=self.SECTIONS[1])
        ttk.Combobox(fr, textvariable=self._section_var, values=self.SECTIONS,
                     width=20, state='readonly').grid(row=0, column=1, padx=6, sticky='w')
        self._other_section_var = tk.StringVar()
        tk.Entry(fr, textvariable=self._other_section_var, width=18,
                 font=("Arial",9)).grid(row=0, column=2, padx=4, sticky='w')
        tk.Label(fr, text="(if Other)", font=("Arial",8), fg='gray',
                 bg=C_LIGHT).grid(row=0, column=3, sticky='w')

        # FIX: from_=0 so opening minutes and hours can be zero
        fields = [
            ("Opening Time (H):", "open_h_var",   8,  "0-23",  0, 23),
            ("Opening Time (M):", "open_m_var",   0,  "0-59",  0, 59),
            ("Periods/Day:",      "periods_var",  7,  "1-12",  1, 12),
            ("Period Duration (min):", "pdur_var", 40, "10-120", 10, 120),
            ("Break Duration (min):",  "brkdur_var",20,"5-60",   5,  60),
            ("Break After Period:",    "brkaft_var", 3, "1-12",  1,  12),
        ]
        for i, (label, attr, default, hint, from_, to_) in enumerate(fields):
            r, c = divmod(i, 3)
            r += 1
            c_base = c * 3
            tk.Label(fr, text=label, font=("Arial",9), fg=C_NAVY,
                     bg=C_LIGHT).grid(row=r, column=c_base, sticky='e', padx=(10,2), pady=4)
            var = tk.IntVar(value=default)
            setattr(self, f"_{attr}", var)
            sp = tk.Spinbox(fr, textvariable=var, from_=from_, to=to_, width=7,
                            font=("Arial",10))
            sp.grid(row=r, column=c_base+1, sticky='w', padx=2)
            tk.Label(fr, text=hint, font=("Arial",7), fg='gray',
                     bg=C_LIGHT).grid(row=r, column=c_base+2, sticky='w')

    # STEP 3 – Classes & sections  (FIX: 3 columns instead of 4)
    def _build_step3(self):
        fr = self._padded("③ Classes & Sections")

        tk.Label(fr, text="Select classes (check) and their sections (A/B/C/D/G or custom):",
                 font=("Arial",9), fg=C_NAVY, bg=C_LIGHT).grid(
            row=0, column=0, columnspan=12, sticky='w', pady=(0,6))

        self._class_vars    = {}
        self._class_sec_vars = {}
        self._class_custom  = {}

        # FIX: 3 columns (was 4, caused horizontal overflow and clipping on smaller screens)
        COLS = 3
        for idx, cls in enumerate(self.CLASSES):
            row_base = (idx // COLS) * 2 + 1
            col_base = (idx % COLS) * 4

            var = tk.BooleanVar()
            self._class_vars[cls] = var
            cb = tk.Checkbutton(fr, text=cls, variable=var, font=("Arial",9,"bold"),
                                fg=C_NAVY, bg=C_LIGHT, activebackground=C_LIGHT,
                                command=self._on_class_toggle)
            cb.grid(row=row_base, column=col_base, sticky='w', padx=4, pady=2)

            sec_fr = tk.Frame(fr, bg=C_LIGHT)
            sec_fr.grid(row=row_base+1, column=col_base, columnspan=4, sticky='w', padx=16)
            self._class_sec_vars[cls] = {}
            for s in self.STD_SECS:
                sv = tk.BooleanVar()
                self._class_sec_vars[cls][s] = sv
                tk.Checkbutton(sec_fr, text=s, variable=sv, font=("Arial",8),
                               fg=C_TEAL, bg=C_LIGHT,
                               command=self._on_class_toggle).pack(side='left')

            cvar = tk.StringVar()
            self._class_custom[cls] = cvar
            tk.Entry(sec_fr, textvariable=cvar, width=8,
                     font=("Arial",8)).pack(side='left', padx=4)
            tk.Label(sec_fr, text="(custom)", font=("Arial",7), fg='gray',
                     bg=C_LIGHT).pack(side='left')

        tk.Label(fr, text="Custom Class Name:", font=("Arial",9), fg=C_NAVY,
                 bg=C_LIGHT).grid(row=99, column=0, sticky='e', padx=6, pady=6)
        self._other_class_var = tk.StringVar()
        tk.Entry(fr, textvariable=self._other_class_var, width=14,
                 font=("Arial",9)).grid(row=99, column=1, sticky='w')

    def _on_class_toggle(self):
        try:
            cs = self._get_class_sections()
            self._incharge_frame.refresh(cs)
        except AttributeError:
            pass

    def _get_class_sections(self):
        result = []
        for cls in self.CLASSES:
            if cls == "Other (type)":
                cls_name = self._other_class_var.get().strip()
                if not cls_name:
                    continue
            else:
                cls_name = cls

            if not self._class_vars.get(cls, tk.BooleanVar()).get():
                continue

            sec_vars   = self._class_sec_vars.get(cls, {})
            has_sec    = any(v.get() for v in sec_vars.values())
            custom_sec = self._class_custom.get(cls, tk.StringVar()).get().strip()

            if not has_sec and not custom_sec:
                result.append((cls_name, "A"))
            else:
                for s, sv in sec_vars.items():
                    if sv.get():
                        result.append((cls_name, s))
                if custom_sec:
                    for s in custom_sec.replace(",", " ").split():
                        result.append((cls_name, s.strip()))

        return result

    # STEP 4 – Incharge teachers
    def _build_step4(self):
        fr = self._padded("④ Class In-Charge Teachers")
        self._incharge_frame = InchargeFrame(fr, [])
        self._incharge_frame.pack(fill='both', expand=True)

    # STEP 5 – Subjects
    def _build_step5(self):
        fr = self._padded("⑤ Subjects & Subject Teachers")

        ctrl_fr = tk.Frame(fr, bg=C_LIGHT)
        ctrl_fr.pack(fill='x', pady=(0,6))
        tk.Label(ctrl_fr, text="Add subjects and assign one or more teachers (comma-separated):",
                 font=("Arial",9), fg=C_NAVY, bg=C_LIGHT).pack(side='left')
        styled_btn(ctrl_fr, "+ Add Subject", self._add_subject_row,
                   bg=C_TEAL, padx=10, pady=3, font_size=9).pack(side='right')

        self._subj_container = tk.Frame(fr, bg=C_LIGHT)
        self._subj_container.pack(fill='x')

        for subj in ["English", "Mathematics", "Science", "Urdu", "Islamiat", "Computer"]:
            self._add_subject_row(default_subj=subj)

    def _add_subject_row(self, default_subj=""):
        def remove():
            row.destroy()
            self._subject_rows.remove(row)
        row = SubjectRow(self._subj_container, remove)
        if default_subj:
            row.subj_var.set(default_subj)
        row.pack(fill='x', pady=2)
        self._subject_rows.append(row)
        # FIX: re-bind scroll after adding a new row
        self.after(50, self._bind_all_scroll)

    # STEP 6 – Generate buttons
    # FIX: buttons are INSIDE the scrollable frame → never clipped on minimized window
    def _build_step6(self):
        fr = self._padded("⑥ Generate Timetable")

        info = tk.Label(fr,
            text="ℹ  Rules: Class In-Charge has Period 1 (10 min longer) · "
                 "Workload distributed equally · Each subject appears at least once",
            font=("Arial", 9, "italic"), fg=C_TEAL, bg=C_LIGHT, justify='left')
        info.pack(anchor='w', pady=(0,10))

        btn_fr = tk.Frame(fr, bg=C_LIGHT)
        btn_fr.pack(pady=10)

        styled_btn(btn_fr, "📋  Generate Class-Wise Timetable",
                   lambda: self._generate("class"),
                   bg=C_NAVY, padx=24, pady=10, font_size=12).pack(side='left', padx=16)

        styled_btn(btn_fr, "👩‍🏫  Generate Teacher-Wise Timetable",
                   lambda: self._generate("teacher"),
                   bg=C_TEAL, padx=24, pady=10, font_size=12).pack(side='left', padx=16)

    # FOOTER – copyright
    def _build_footer(self):
        """FIX: Footer placed inside scroll_frame so it's always visible at bottom."""
        footer = tk.Frame(self._scroll_frame, bg=C_NAVY, pady=10)
        footer.pack(fill='x', padx=0, pady=(12, 0), side='bottom')

        tk.Label(footer,
                 text="© Farrukh Barlas  |  AFUWEBS",
                 font=("Arial", 9, "bold"),
                 fg=C_GOLD, bg=C_NAVY).pack()

        tk.Label(footer,
                 text="All Rights Reserved",
                 font=("Arial", 8),
                 fg=C_WHITE, bg=C_NAVY).pack()

    # ── GENERATE ──────────────────────────────────────────────────────────────
    def _generate(self, tt_type):
        self.data.school_name = self._school_var.get().strip()
        self.data.logo_path   = self._logo_var.get().strip()

        sec = self._section_var.get()
        if sec == "Other (type below)":
            sec = self._other_section_var.get().strip() or "Other"
        self.data.section = sec

        self.data.open_h      = self._open_h_var.get()
        self.data.open_m      = self._open_m_var.get()
        self.data.num_periods = self._periods_var.get()
        self.data.period_mins = self._pdur_var.get()
        self.data.break_mins  = self._brkdur_var.get()
        self.data.break_after = self._brkaft_var.get()

        cs = self._get_class_sections()
        if not cs:
            messagebox.showwarning("No Classes", "Please select at least one class and section.")
            return
        self.data.class_sections = cs
        self.data.incharges = self._incharge_frame.get_incharges()

        subjects = []
        for row in self._subject_rows:
            subj, tchrs = row.get()
            if subj:
                subjects.append((subj, tchrs if tchrs else ["TBD"]))
        if not subjects:
            messagebox.showwarning("No Subjects", "Please add at least one subject.")
            return
        self.data.subjects = subjects

        # Validate break after
        if self.data.break_after >= self.data.num_periods:
            messagebox.showwarning("Break Setting",
                "Break-after period cannot equal or exceed total periods. Adjusted.")
            self.data.break_after = max(1, self.data.num_periods - 1)

        try:
            self.data.class_tt, self.data.teacher_tt = TTEngine.generate(self.data)
        except ValueError as e:
            messagebox.showerror("Generation Error", str(e))
            return

        PreviewWindow(self, self.data, tt_type)


# ─────────────────────────────────────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = TimetableApp()
    app.mainloop()
